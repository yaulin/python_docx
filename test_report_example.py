#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sat Nov 11 13:41:36 2023

@author: yaroslav
"""

import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml import OxmlElement, ns

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
sns.set(font_scale=1.5)


from datetime import date


# change the parameters to generate test report
sn = 10000  # serial number of the device
device_type = 'Power' # device type Standard or Power
wl = '785' # wavelength or list of wavelengths
operator = "Yaroslav Aulin" #operator name
f = "polystyrene" # .tsv data file

# functions to add page number

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    
    
    
# raman spectrum class
class raman_spectrum:
    def __init__(self,fname):
        self.fname = fname
        self.fname0 = self.fname.split('/')[-1]
        self.df = pd.read_csv(fname + '.tsv', sep='\t', header =7)
        self.df.rename(columns={self.df.columns[0]:'Raman Shift, cm-1', self.df.columns[1]:'Intensity, %'}, inplace = True)
        self.df.replace([np.inf, -np.inf], np.nan, inplace=True)
        self.df.dropna(inplace = True)
        self.df["Normalized Intensity"] = (self.df["Intensity, %"] - (min0:= self.df["Intensity, %"].min())) / (self.df["Intensity, %"].max()- min0 )
    
    def normalize1(self,min0):
        self.df["Normalized Intensity"] = (self.df["Intensity, %"] - min0) / (self.df["Intensity, %"].max()- min0 )
    
    
    def plot(self, color = "red", label = ""):
        sns.lineplot(self.df, x = "Raman Shift, cm-1", y = "Intensity, %", color = color, label = label)
        plt.xlabel(r'Raman Shift, $cm^{-1}$')
        
    def plot_n(self, color, label, marker = None):
        sns.lineplot(self.df, x = "Raman Shift, cm-1", y = "Normalized Intensity", marker = marker, color = color, label = label)
        plt.xlabel(r'Raman Shift, $cm^{-1}$') 
        
    def get_max(self):
        return self.df["Intensity, %"].max()
    def get_peak_position(self):
        idx = self.df["Intensity, %"].idxmax()
        return self.df["Raman Shift, cm-1"][idx]
    
    


    

# create a .docx document
document = docx.Document()



# generate header and footer

section = document.sections[0]
header = section.header
footer = section.footer


p = header.paragraphs[0]
r = p.add_run()
r.add_picture("lightnovo-logo-red-current.png", width = Inches(2))
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = footer.paragraphs[0]
r = p.add_run()
r.text = "TEST REPORT\tminiRaman SN" + str(sn) + "\t"

add_page_number(document.sections[0].footer.paragraphs[0].add_run())



# add heading
document.add_heading('Certificate of Quality', 0)

# add equipment type and serial number 
p = document.add_paragraph()
r = p.add_run()
r.add_text("Equipment: miniRaman " + device_type + ' ' + wl + ' nm')

p = document.add_paragraph()
r = p.add_run()
r.add_text("Serial number: #" + str(sn) )


# add subtitle
document.add_heading('Polystyrene spectrum', 1)

# load data and plot a graph
s = raman_spectrum(f)
plt.figure(figsize = (10,5))
s.plot()
plt.savefig("polystyrene.png",dpi = 300, bbox_inches='tight')


# add figure to docx
p = document.add_paragraph()
r = p.add_run()
r.add_picture("polystyrene.png", width = Inches(5))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# add list
p = document.add_paragraph()
p.style = 'List Bullet'
r = p.add_run()
r.add_text("Max intensity: {:10.1f}".format(s.get_max()) + "%")

p = document.add_paragraph()
p.style = 'List Bullet'
r = p.add_run()
r.add_text("Peak position: {:10.1f} ".format(s.get_peak_position()) + r"cm^-1")

p = document.add_paragraph()
p.style = 'List Bullet'
r = p.add_run()
r.add_text("Test result: ")
r = p.add_run()
r.add_text("pass")
r.font.color.rgb = RGBColor.from_string('056608')

#add some empty space
p = document.add_paragraph()
r = p.add_run()
for i in range(0,5):
    r.add_break()

# add date
p = document.add_paragraph()
r = p.add_run()
today = date.today()
r.add_text("Date: " + str(today))


# add operator name
p = document.add_paragraph()
r = p.add_run()
today = date.today()
r.add_text("Operator: " + operator)


# save document
document.save("test_report.docx")